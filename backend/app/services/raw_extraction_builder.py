import os
import re

import fitz
import pdfplumber
from docx import Document


class RawExtractionBuilder:

    @staticmethod
    def build(file_path):
        extension = os.path.splitext(file_path)[1].lower()

        if extension == ".pdf":
            return RawExtractionBuilder._build_pdf(file_path)

        if extension == ".docx":
            return RawExtractionBuilder._build_docx(file_path)

        raise ValueError("Unsupported report file type")

    @staticmethod
    def _build_pdf(file_path):
        doc = fitz.open(file_path)
        text_blocks = []
        images = []
        full_text_parts = []

        for page_index in range(len(doc)):
            page = doc[page_index]
            page_number = page_index + 1
            width = page.rect.width
            height = page.rect.height
            page_text = page.get_text() or ""
            full_text_parts.append(page_text)

            for order, block in enumerate(page.get_text("blocks") or []):
                text = RawExtractionBuilder._clean_text(block[4])
                if not text:
                    continue

                location = RawExtractionBuilder._location(
                    page_number,
                    width,
                    height,
                    block[0],
                    block[1],
                    block[2],
                    block[3]
                )
                text_blocks.append({
                    "text": text,
                    "page_number": page_number,
                    "order": order,
                    "location": location,
                    "zone": RawExtractionBuilder._zone_from_location(location),
                })

            for image_index, image in enumerate(page.get_images(full=True) or []):
                rects = []
                try:
                    rects = page.get_image_rects(image[0])
                except Exception:
                    rects = []

                location = None
                if rects:
                    rect = rects[0]
                    location = RawExtractionBuilder._location(
                        page_number,
                        width,
                        height,
                        rect.x0,
                        rect.y0,
                        rect.x1,
                        rect.y1
                    )

                images.append({
                    "page_number": page_number,
                    "image_index": image_index,
                    "xref": image[0],
                    "location": location,
                    "zone": (
                        RawExtractionBuilder._zone_from_location(location)
                        if location
                        else None
                    ),
                    "caption": None,
                })

        tables = RawExtractionBuilder._extract_pdf_tables(file_path)
        captions = RawExtractionBuilder._extract_captions(
            "\n".join(full_text_parts)
        )
        location_captions = RawExtractionBuilder._extract_captions_below_images(
            images,
            text_blocks
        )
        captions.extend(
            caption
            for caption in location_captions
            if caption not in captions
        )
        RawExtractionBuilder._attach_captions(images, captions)

        return {
            "model_type": "raw_extraction",
            "file_type": "pdf",
            "page_count": len(doc),
            "text": "\n".join(full_text_parts),
            "text_blocks": text_blocks,
            "paragraphs": [],
            "tables": tables,
            "images": images,
        }

    @staticmethod
    def _build_docx(file_path):
        document = Document(file_path)
        paragraphs = []
        text_parts = []
        header_paragraphs = []

        for section in document.sections:

            for paragraph in section.header.paragraphs:

                text = RawExtractionBuilder._clean_text(paragraph.text)

                if text:

                    header_paragraphs.append(text)

        footer_paragraphs = []

        for section in document.sections:

            for paragraph in section.footer.paragraphs:

                text = RawExtractionBuilder._clean_text(paragraph.text)

                if text:
                    footer_paragraphs.append(text)
        
        all_paragraphs = []

        for text in header_paragraphs:

            all_paragraphs.append(
                {
                    "text": text,
                    "source": "header",
                }
            )

        for text in footer_paragraphs:

            all_paragraphs.append(
                {
                    "text": text,
                    "source": "footer",
                }
            )

        for paragraph in document.paragraphs:

            text = RawExtractionBuilder._clean_text(paragraph.text)

            if text:

                all_paragraphs.append(
                    {
                        "text": text,
                        "source": "body",
                    }
                )

        for p in document.sections[0].footer.paragraphs:
            print(repr(p.text))

        for index, paragraph in enumerate(all_paragraphs):

            text = paragraph["text"]

            source = paragraph["source"]

            if not text:
                continue

            text_parts.append(text)

            if source == "header":

                zone = "top_center"

            zone = None

            if source == "header":

                zone = "top_center"

            elif source == "footer":

                lower = text.lower()

                if "principal" in lower:

                    zone = "bottom_right"

                elif (
                    "faculty" in lower
                    or "club coordinator" in lower
                    or "coordinator" in lower
                ):

                    zone = "bottom_left"

                else:

                    zone = "bottom_left"

            else:

                if index <= 2:

                    zone = "top_center"

            if source == "header":

                location = RawExtractionBuilder._location(
                    1, 595, 842,
                    40, 30,
                    555, 110,
                )

            elif source == "footer":

                location = RawExtractionBuilder._location(
                    1, 595, 842,
                    40, 760,
                    555, 790,
                )

            else:

                body_index = index - len(header_paragraphs)

                y = 130 + body_index * 25

                location = RawExtractionBuilder._location(
                    1, 595, 842,
                    40,
                    y,
                    555,
                    y + 20,
                )

            paragraphs.append({
                "text": text,
                "order": index,
                "page_number": 1,
                "zone": zone,
                "location": location,
            })

        tables = []
        for table_index, table in enumerate(document.tables):
            rows = []
            for row_index, row in enumerate(table.rows):
                cells = [
                    RawExtractionBuilder._clean_text(cell.text)
                    for cell in row.cells
                ]
                rows.append({
                    "row_order": row_index,
                    "cells": cells,
                    "label": RawExtractionBuilder._first_non_empty(cells),
                    "value": RawExtractionBuilder._second_non_empty(cells),

                    "location": {
                        "page_number": 1,
                    },
                })

            tables.append({
                "table_index": table_index,
                "page_number": None,
                "rows": rows,
            })

        captions = RawExtractionBuilder._extract_captions("\n".join(text_parts))

        if not captions:

            for paragraph in paragraphs:

                text = paragraph.get("text", "")

                zone = paragraph.get("zone")

                lower = text.lower()

                if (
                    zone is None
                    and 2 <= len(text.split()) <= 10
                    and "report" not in lower
                    and "institute" not in lower
                    and "coimbatore" not in lower
                    and "faculty" not in lower
                    and "principal" not in lower
                ):
                    captions.append(text)

        images = []
        try:
            image_count = len(document.inline_shapes)
        except Exception:
            image_count = 0

        for index in range(image_count):

            if image_count >= 3:

                if index == 0:

                    location = RawExtractionBuilder._location(
                        1,595,842,
                        20,20,90,90
                    )

                elif index == 1:

                    location = RawExtractionBuilder._location(
                        1,595,842,
                        505,20,575,90
                    )

                else:

                    location = RawExtractionBuilder._location(
                        1,
                        595,
                        842,
                        120,
                        330,
                        480,
                        520,
                    )

            elif image_count == 2:

                if header_paragraphs:

                    if index == 0:

                        location = RawExtractionBuilder._location(
                            1,595,842,
                            20,20,90,90
                        )

                    else:

                        location = RawExtractionBuilder._location(
                            1,
                            595,
                            842,
                            120,
                            330,
                            480,
                            520,
                        )

                else:

                    location = RawExtractionBuilder._location(
                        1,
                        595,
                        842,
                        120,
                        330,
                        480,
                        520,
                    )

            else:

                if header_paragraphs:

                    location = RawExtractionBuilder._location(
                        1,
                        595,
                        842,
                        120,
                        330,
                        480,
                        520,
                    )

                else:

                    location = RawExtractionBuilder._location(
                        1,
                        595,
                        842,
                        20,
                        20,
                        90,
                        90,
                    )

            zone = RawExtractionBuilder._zone_from_location(location)

            images.append(
                {
                    "page_number": 1,
                    "image_index": index,
                    "location": location,
                    "zone": zone,
                    "caption": captions[0] if zone == "center" and captions else None,
                    "is_header": zone.startswith("top"),
                }
            )

        return {
            "model_type": "raw_extraction",
            "file_type": "docx",
            "page_count": None,
            "text": "\n".join(text_parts),
            "text_blocks": [],
            "paragraphs": paragraphs,
            "tables": tables,
            "images": images,
        }

    @staticmethod
    def _extract_pdf_tables(file_path):
        tables = []

        with pdfplumber.open(file_path) as pdf:

            for page_index, page in enumerate(pdf.pages):

                extracted_tables = page.find_tables()

                for table_index, table in enumerate(extracted_tables):

                    extracted = table.extract()

                    rows = []

                    for row_index, row in enumerate(extracted):

                        cells = [
                            RawExtractionBuilder._clean_text(cell or "")
                            for cell in row
                        ]

                        row_location = None

                        try:

                            if row_index < len(table.rows):

                                bbox = table.rows[row_index].bbox

                                x0, y0, x1, y1 = bbox

                                row_location = {
                                    "page_number": page_index + 1,
                                    "page_width": page.width,
                                    "page_height": page.height,
                                    "x0": x0,
                                    "y0": y0,
                                    "x1": x1,
                                    "y1": y1,
                                    "center_x": (x0 + x1) / 2,
                                    "center_y": (y0 + y1) / 2,
                                }

                        except Exception:
                            pass

                        rows.append(
                            {
                                "row_order": row_index,
                                "cells": cells,
                                "label": RawExtractionBuilder._first_non_empty(cells),
                                "value": RawExtractionBuilder._second_non_empty(cells),
                                "location": (
                                    row_location
                                    if row_location
                                    else {
                                        "page_number": page_index + 1,
                                        "page_width": page.width,
                                        "page_height": page.height,
                                        "x0": 0,
                                        "y0": 0,
                                        "x1": page.width,
                                        "y1": 0,
                                        "center_x": page.width / 2,
                                        "center_y": 0,
                                    }
                                ),
                            }
                        )

                    tables.append(
                        {
                            "table_index": table_index,
                            "page_number": page_index + 1,
                            "rows": rows,
                        }
                    )

        return tables

    @staticmethod
    def _location(page_number, page_width, page_height, x0, y0, x1, y1):
        return {
            "page_number": page_number,
            "page_width": page_width,
            "page_height": page_height,
            "x0": x0,
            "y0": y0,
            "x1": x1,
            "y1": y1,
            "center_x": (x0 + x1) / 2,
            "center_y": (y0 + y1) / 2,
        }

    @staticmethod
    def _zone_from_location(location):
        page_width = location.get("page_width") or 1
        page_height = location.get("page_height") or 1
        center_x = location.get("center_x", 0) / page_width
        center_y = location.get("center_y", 0) / page_height

        if center_y < 0.25:
            if center_x < 0.33:
                return "top_left"
            if center_x > 0.67:
                return "top_right"
            return "top_center"

        if center_y > 0.75:
            return "bottom_left" if center_x < 0.5 else "bottom_right"

        return "center"

    @staticmethod
    def _extract_captions(text):
        captions = []
        pattern = re.compile(
            r"(?im)^\s*(?:caption(?:\s+for\s+the\s+image)?|fig(?:ure)?|image|photo)"
            r"\s*[\d.:-]*\s*(.+?)\s*$"
        )

        for match in pattern.finditer(text or ""):
            caption = RawExtractionBuilder._clean_text(match.group(1))
            if caption and caption not in captions:
                captions.append(caption)

        return captions

    @staticmethod
    def _attach_captions(images, captions):
        for index, caption in enumerate(captions):
            if index < len(images):
                images[index]["caption"] = caption

    @staticmethod
    def _extract_captions_below_images(images, text_blocks):
        captions = []
        signature_pattern = re.compile(
            r"\b(?:faculty|club)\s+co-?ordinator\b"
            r"|\bprincipal\b"
            r"|\bsecretary\b"
            r"|\bjoint secretary\b",
            re.IGNORECASE
        )

        for image in images:
            image_location = image.get("location")
            if not image_location:
                continue
            page_height = image_location.get("page_height") or 1
            if image.get("zone") in {"top_left", "top_center", "top_right"}:
                continue
            if (image_location.get("center_y", 0) / page_height) < 0.2:
                continue

            image_page = image_location.get("page_number")
            image_x0 = image_location.get("x0", 0)
            image_x1 = image_location.get("x1", 0)
            image_y1 = image_location.get("y1", 0)
            image_width = max(image_x1 - image_x0, 1)
            nearby_blocks = []

            for block in text_blocks:
                location = block.get("location") or {}
                if location.get("page_number") != image_page:
                    continue

                y0 = location.get("y0", 0)
                if y0 < image_y1 or y0 - image_y1 > 70:
                    continue

                overlap = min(image_x1, location.get("x1", 0)) - max(
                    image_x0,
                    location.get("x0", 0)
                )
                if overlap < image_width * 0.15:
                    continue

                nearby_blocks.append(block)

            nearby_blocks.sort(
                key=lambda block: (
                    (block.get("location") or {}).get("y0", 0),
                    (block.get("location") or {}).get("x0", 0)
                )
            )

            caption_parts = []
            for block in nearby_blocks:
                text = RawExtractionBuilder._clean_text(block.get("text", ""))
                if not text:
                    continue

                if signature_pattern.search(text):
                    text = signature_pattern.split(text, maxsplit=1)[0].strip()

                if text:
                    caption_parts.append(text)

            caption = RawExtractionBuilder._clean_text(" ".join(caption_parts))
            if caption and caption not in captions:
                captions.append(caption)

        return captions

    @staticmethod
    def _first_non_empty(cells):
        for cell in cells:
            if cell:
                return cell
        return None

    @staticmethod
    def _second_non_empty(cells):
        values = [cell for cell in cells if cell]
        return values[1] if len(values) > 1 else None

    @staticmethod
    def _clean_text(value):
        value = re.sub(r"[\u200b\u200c\u200d\ufeff]", " ", str(value or ""))
        return re.sub(r"\s+", " ", value).strip()
