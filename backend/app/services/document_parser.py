from __future__ import annotations

import re
from typing import Optional

import fitz
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

try:
    from app.services.layout_extractor import LayoutExtractor
    _HAS_LAYOUT_EXTRACTOR = True
except ImportError:
    _HAS_LAYOUT_EXTRACTOR = False

_ZONE_DEFS = {
    "top_left":     {"x": (0.00, 0.33), "y": (0.00, 0.25)},
    "top_center":   {"x": (0.33, 0.67), "y": (0.00, 0.25)},
    "top_right":    {"x": (0.67, 1.00), "y": (0.00, 0.25)},
    "center":       {"x": (0.25, 0.75), "y": (0.25, 0.75)},
    "bottom_left":  {"x": (0.00, 0.50), "y": (0.75, 1.00)},
    "bottom_right": {"x": (0.50, 1.00), "y": (0.75, 1.00)},
}


def _classify_zone(rel_cx: float, rel_cy: float) -> Optional[str]:
    for zone, bounds in _ZONE_DEFS.items():
        if (bounds["x"][0] <= rel_cx <= bounds["x"][1]
                and bounds["y"][0] <= rel_cy <= bounds["y"][1]):
            return zone
    return None


def _clean(value) -> str:
    value = re.sub(r"[\u200b\u200c\u200d\ufeff]", " ", str(value or ""))
    return re.sub(r"\s+", " ", value).strip()

def _pdf_zone(x0, y0, x1, y1, page_width, page_height) -> Optional[str]:
    cx = ((x0 + x1) / 2) / max(page_width, 1)
    cy = ((y0 + y1) / 2) / max(page_height, 1)
    return _classify_zone(cx, cy)


def _location_box(x0, y0, x1, y1, page_width, page_height, page_number):
    return {
        "page_number": page_number,
        "page_width": page_width,
        "page_height": page_height,
        "x0": x0,
        "y0": y0,
        "x1": x1,
        "y1": y1,
        "center_x": (x0 + x1) / 2,
        "center_y": (y0 + y1) / 2,
    }

_ALIGN_MAP = {
    WD_ALIGN_PARAGRAPH.LEFT: "left",
    WD_ALIGN_PARAGRAPH.CENTER: "center",
    WD_ALIGN_PARAGRAPH.RIGHT: "right",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
}

def _docx_alignment(paragraph) -> Optional[str]:
    return _ALIGN_MAP.get(paragraph.alignment)

def _paragraph_is_bold(paragraph) -> bool:
    runs = [r for r in paragraph.runs if r.text.strip()]
    if not runs:
        return False
    return all(r.bold for r in runs)

_A4_LINES_PER_PAGE = 45  # conservative estimate for a typical A4 document


def _estimate_docx_page_count(paragraphs: list) -> int:

    line_count = sum(1 for p in paragraphs if p.get("text", "").strip())
    return max(1, round(line_count / _A4_LINES_PER_PAGE + 0.5))

_A4_WIDTH = 595.0
_A4_HEIGHT = 842.0

_LINE_HEIGHT = 14.0
_CHAR_WIDTH = 7.0
_LEFT_MARGIN = 72.0
_RIGHT_MARGIN = _A4_WIDTH - 72.0


def _docx_pseudo_layout(
    doc_paragraphs: list[dict],
    header_lines: list[str],
    footer_lines: list[str],
) -> list[dict]:

    pages: dict[int, list] = {}

    header_set = set(header_lines)
    footer_set = set(footer_lines)

    h_y = 10.0

    for line in header_lines:

        if not line:
            continue

        w = min(
            len(line) * _CHAR_WIDTH,
            _RIGHT_MARGIN - _LEFT_MARGIN,
        )

        cx = _A4_WIDTH / 2

        x0 = cx - w / 2
        x1 = cx + w / 2

        y0 = h_y
        y1 = h_y + _LINE_HEIGHT

        location = _location_box(
            x0,
            y0,
            x1,
            y1,
            _A4_WIDTH,
            _A4_HEIGHT,
            1,
        )

        pages.setdefault(1, []).append(
            {
                "text": line,
                "x0": x0,
                "y0": y0,
                "x1": x1,
                "y1": y1,
                "zone": _pdf_zone(
                    x0,
                    y0,
                    x1,
                    y1,
                    _A4_WIDTH,
                    _A4_HEIGHT,
                ),
                "location": location,
            }
        )

        h_y += _LINE_HEIGHT

    body_y = max(
        h_y + _LINE_HEIGHT,
        120.0,
    )

    usable_page_height = _A4_HEIGHT - 200.0

    for para in doc_paragraphs:

        text = para.get("text", "").strip()

        if (
            not text
            or text in header_set
            or text in footer_set
        ):
            continue

        alignment = para.get("alignment")
        style = para.get("style", "")

        w = min(
            len(text) * _CHAR_WIDTH,
            _RIGHT_MARGIN - _LEFT_MARGIN,
        )

        if (
            alignment == "center"
            or "heading" in style.lower()
            or "title" in style.lower()
        ):

            cx = _A4_WIDTH / 2

        elif alignment == "right":

            cx = _RIGHT_MARGIN - w / 2

        else:

            cx = _LEFT_MARGIN + w / 2

        x0 = max(
            _LEFT_MARGIN,
            cx - w / 2,
        )

        x1 = min(
            _RIGHT_MARGIN,
            cx + w / 2,
        )

        page_num = (
            1
            + int(
                (body_y - 120.0)
                / usable_page_height
            )
        )

        local_y = (
            body_y
            - (page_num - 1) * usable_page_height
            + 120.0
        )

        y0 = local_y
        y1 = local_y + _LINE_HEIGHT

        location = _location_box(
            x0,
            y0,
            x1,
            y1,
            _A4_WIDTH,
            _A4_HEIGHT,
            page_num,
        )

        pages.setdefault(page_num, []).append(
            {
                "text": text,
                "x0": x0,
                "y0": y0,
                "x1": x1,
                "y1": y1,
                "zone": _pdf_zone(
                    x0,
                    y0,
                    x1,
                    y1,
                    _A4_WIDTH,
                    _A4_HEIGHT,
                ),
                "location": location,
            }
        )

        body_y += _LINE_HEIGHT

    last_page = max(
        pages.keys(),
        default=1,
    )

    f_y = _A4_HEIGHT - 80.0

    for line in footer_lines:

        if not line:
            continue

        w = min(
            len(line) * _CHAR_WIDTH,
            _RIGHT_MARGIN - _LEFT_MARGIN,
        )

        cx = _A4_WIDTH / 2

        x0 = cx - w / 2
        x1 = cx + w / 2

        y0 = f_y
        y1 = f_y + _LINE_HEIGHT

        location = _location_box(
            x0,
            y0,
            x1,
            y1,
            _A4_WIDTH,
            _A4_HEIGHT,
            last_page,
        )

        pages.setdefault(last_page, []).append(
            {
                "text": line,
                "x0": x0,
                "y0": y0,
                "x1": x1,
                "y1": y1,
                "zone": _pdf_zone(
                    x0,
                    y0,
                    x1,
                    y1,
                    _A4_WIDTH,
                    _A4_HEIGHT,
                ),
                "location": location,
            }
        )

        f_y += _LINE_HEIGHT

    return [
        {
            "page_number": page_num,
            "width": _A4_WIDTH,
            "height": _A4_HEIGHT,
            "text_blocks": blocks,
        }
        for page_num, blocks in sorted(
            pages.items()
        )
    ]

def _docx_header_footer(doc: Document) -> tuple[list[str], list[str]]:

    header_lines: list[str] = []
    footer_lines: list[str] = []

    for section in doc.sections:
        try:
            hdr = section.header
            if hdr and not hdr.is_linked_to_previous:
                for para in hdr.paragraphs:
                    line = _clean(para.text)
                    if line and line not in header_lines:
                        header_lines.append(line)
        except Exception:
            pass

        try:
            ftr = section.footer
            if ftr and not ftr.is_linked_to_previous:
                for para in ftr.paragraphs:
                    line = _clean(para.text)
                    if line and line not in footer_lines:
                        footer_lines.append(line)
        except Exception:
            pass

    return header_lines, footer_lines

def _docx_images(doc: Document) -> list[dict]:
    images = []

    idx = 0
    header_idx = 0

    for section in doc.sections:
        for para in section.header.paragraphs:
            for run in para.runs:

                drawing = run._element.find(qn("w:drawing"))

                if drawing is None:
                    drawing = run._element.find(qn("w:pict"))

                if drawing is None:
                    continue

                if header_idx == 0:
                    zone = "top_left"
                else:
                    zone = "top_right"

                images.append({
                    "index": idx,
                    "page_number": 1,
                    "zone": zone,
                    "is_header": True,
                    "width": None,
                    "height": None,
                    "location": None,
                })

                idx += 1
                header_idx += 1

    for para in doc.paragraphs:
        for run in para.runs:

            drawing = run._element.find(qn("w:drawing"))

            if drawing is None:
                drawing = run._element.find(qn("w:pict"))

            if drawing is None:
                continue

            images.append({
                "index": idx,
                "page_number": 1,
                "zone": None,
                "is_header": False,
                "width": None,
                "height": None,
                "location": None,
            })

            idx += 1

    return images

def _docx_page_count_from_props(doc: Document) -> Optional[int]:
    try:
        pages = doc.core_properties.words  # often None in templates
        if pages and isinstance(pages, int) and pages > 0:
            return pages
    except Exception:
        pass
    return None

class DocumentParser:
 
    @staticmethod
    def parse_pdf(file_path: str) -> dict:
        doc = fitz.open(file_path)

        full_text_parts: list[str] = []
        paragraphs: list[dict] = []
        layout: list[dict] = []
        images: list[dict] = []
        image_idx = 0

        for page_obj in doc:
            page_num = page_obj.number + 1
            pw = page_obj.rect.width or _A4_WIDTH
            ph = page_obj.rect.height or _A4_HEIGHT

            page_text_blocks: list[dict] = []
            raw_dict = page_obj.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)
            order_idx = 0

            for block in raw_dict.get("blocks", []):
                if block.get("type") != 0:  # 0 = text
                    continue
                bx0, by0, bx1, by1 = block["bbox"]
                block_text_lines = []

                for line in block.get("lines", []):
                    line_text = " ".join(
                        span.get("text", "") for span in line.get("spans", [])
                    ).strip()
                    if line_text:
                        block_text_lines.append(line_text)

                block_text = "\n".join(block_text_lines)
                if not block_text.strip():
                    continue

                zone = _pdf_zone(bx0, by0, bx1, by1, pw, ph)
                loc = _location_box(bx0, by0, bx1, by1, pw, ph, page_num)

                is_bold = False
                alignment = "left"
                first_line = block.get("lines", [{}])[0]
                first_span = first_line.get("spans", [{}])[0]
                font_flags = first_span.get("flags", 0)
                is_bold = bool(font_flags & 2 ** 4)  # bold flag in PyMuPDF

                bcx = (bx0 + bx1) / 2
                if abs(bcx / pw - 0.5) < 0.12:
                    alignment = "center"
                elif bx0 / pw > 0.6:
                    alignment = "right"

                tb = {
                    "text": block_text,
                    "zone": zone,
                    "order": order_idx,
                    "page_number": page_num,
                    "location": loc,
                    "bold": is_bold,
                    "alignment": alignment,
                    "style": "Bold" if is_bold else "Normal",
                }
                page_text_blocks.append(tb)
                paragraphs.append(tb)
                full_text_parts.append(block_text)
                order_idx += 1

            try:
                for img_info in page_obj.get_images(full=True):
                    xref = img_info[0]
                    rects = page_obj.get_image_rects(xref)
                    if rects:
                        r = rects[0]
                        ix0, iy0, ix1, iy1 = r.x0, r.y0, r.x1, r.y1
                    else:
                        ix0, iy0, ix1, iy1 = 0, 0, 0, 0

                    zone = _pdf_zone(ix0, iy0, ix1, iy1, pw, ph) if rects else None
                    loc = _location_box(ix0, iy0, ix1, iy1, pw, ph, page_num) if rects else None
                    is_header = zone in ("top_left", "top_center", "top_right")
                    img_w = (ix1 - ix0) / pw if rects else None
                    img_h = (iy1 - iy0) / ph if rects else None

                    images.append({
                        "index": image_idx,
                        "page_number": page_num,
                        "zone": zone,
                        "is_header": is_header,
                        "width": img_w,
                        "height": img_h,
                        "location": loc,
                    })
                    image_idx += 1
            except Exception:
                pass

            page_layout: dict = {
                "page_number": page_num,
                "width": pw,
                "height": ph,
                "text_blocks": page_text_blocks,
            }

            if _HAS_LAYOUT_EXTRACTOR:
                try:
                    extracted = LayoutExtractor.extract_pdf_layout(doc)
                    if extracted and len(extracted) >= page_num:
                        page_layout = extracted[page_num - 1]
                        page_layout.setdefault("text_blocks", page_text_blocks)
                except Exception:
                    pass

            layout.append(page_layout)

        return {
            "file_type": "pdf",
            "page_count": len(doc),
            "image_count": image_idx,
            "text": "\n".join(full_text_parts),
            "paragraphs": paragraphs,
            "tables": [],          # populated by TemplateAnalyzer
            "header": {
                "text": "",
                "lines": [],
            },
            "footer": {
                "text": "",
                "lines": [],
            },
            "images": images,
            "layout": layout,
        }

    @staticmethod
    def parse_docx(file_path: str) -> dict:
        doc = Document(file_path)

        header_lines, footer_lines = _docx_header_footer(doc)

        para_records: list[dict] = []
        full_text_parts: list[str] = []
        order_idx = 0

        for para in doc.paragraphs:
            text = _clean(para.text)
            if not text:
                order_idx += 1
                continue

            alignment = _docx_alignment(para)
            bold = _paragraph_is_bold(para)
            style_name = para.style.name if para.style else "Normal"

            para_records.append({
                "text": text,
                "style": style_name,
                "bold": bold,
                "alignment": alignment,
                "order": order_idx,
                "page_number": None,   # DOCX has no per-paragraph page info
                "location": None,
                "zone": None,
            })
            full_text_parts.append(text)
            order_idx += 1

        all_text_lines = header_lines + full_text_parts + footer_lines

        images = _docx_images(doc)

        try:
            image_count = len(doc.inline_shapes) or len(images)

        except Exception:
            image_count = len(images)

        header_image_count = sum(
            1
            for img in images
            if img.get("is_header")
)

        page_count = _docx_page_count_from_props(doc) or _estimate_docx_page_count(para_records)

        layout = _docx_pseudo_layout(para_records, header_lines, footer_lines)

        _assign_docx_zones(para_records, layout)

        tables: list[dict] = []
        for t_idx, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                rows.append([
                    _clean(cell.text) for cell in row.cells
                ])
            tables.append({
                "table_index": t_idx,
                "page_number": None,
                "rows": rows,
            })

        return {
            "file_type": "docx",
            "page_count": page_count,
            "paragraph_count": len(para_records),
            "image_count": image_count,
            "text": "\n".join(all_text_lines),
            "paragraphs": para_records,
            "tables": tables,
            "header": {
                "text": "\n".join(header_lines),
                "lines": header_lines,
            },
            "footer": {
                "text": "\n".join(footer_lines),
                "lines": footer_lines,
            },
            "images": images,
            "header_image_count": header_image_count,
            "layout": layout,
        }

def _assign_docx_zones(para_records: list[dict], layout: list[dict]) -> None:
    text_to_zone: dict[str, str] = {}
    for page in layout:
        for block in page.get("text_blocks", []):
            txt = block.get("text", "")
            zone = block.get("zone")
            if txt and zone:
                text_to_zone[txt] = zone

    for para in para_records:
        para["zone"] = text_to_zone.get(para.get("text", ""))

# Claude Generated 2nd time.